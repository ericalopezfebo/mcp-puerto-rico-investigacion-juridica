"""Deterministic local document privacy helpers.

There are deliberately no network imports or outbound requests in this module.
The only supported input is a file inside VELUM_DOCUMENT_ROOT (or the default
~/Documents/VELUM directory). The AI receives only the sanitized text returned
by the MCP tool, never the original file bytes.
"""

from __future__ import annotations

import hashlib
import json
import os
import re
from pathlib import Path
from typing import Any

from bs4 import BeautifulSoup
from pypdf import PdfReader

DEFAULT_ROOT = Path.home() / "Documents" / "VELUM"
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown", ".html", ".htm"}
MAX_FILE_BYTES = 25 * 1024 * 1024
MAX_OUTPUT_CHARS = 60000


def document_root() -> Path:
    configured = os.environ.get("VELUM_DOCUMENT_ROOT", "").strip()
    root = Path(configured).expanduser() if configured else DEFAULT_ROOT
    root.mkdir(parents=True, exist_ok=True)
    return root.resolve()


def safe_path(ruta: str, *, must_exist: bool = True) -> Path:
    if not ruta or not str(ruta).strip():
        raise ValueError("Debes indicar la ruta de un documento local.")

    root = document_root()
    candidate = Path(ruta).expanduser()
    if not candidate.is_absolute():
        candidate = root / candidate
    candidate = candidate.resolve(strict=False)

    try:
        candidate.relative_to(root)
    except ValueError as exc:
        raise ValueError(
            f"Por seguridad, solo se permiten archivos dentro de VELUM_DOCUMENT_ROOT: {root}"
        ) from exc

    if must_exist and not candidate.exists():
        raise FileNotFoundError(f"No existe el documento local: {candidate.name}")
    if candidate.is_dir():
        raise ValueError("La ruta apunta a una carpeta, no a un documento.")
    if candidate.suffix.lower() not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"Tipo no permitido: {candidate.suffix or '[sin extensión]'}. "
            f"Permitidos: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )
    if must_exist and candidate.stat().st_size > MAX_FILE_BYTES:
        raise ValueError(f"El archivo supera el límite local de {MAX_FILE_BYTES // (1024 * 1024)} MB.")
    return candidate


def extract_text(path: Path) -> tuple[str, str]:
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        reader = PdfReader(str(path))
        pages: list[str] = []
        for number, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(f"[página {number}]\n{text}")
        return "\n\n".join(pages), "PDF"

    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("Falta python-docx. Ejecuta: pip install -e .") from exc
        document = Document(str(path))
        blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append(" | ".join(cells))
        return "\n\n".join(blocks), "DOCX"

    raw = path.read_text(encoding="utf-8", errors="replace")
    if suffix in {".html", ".htm"}:
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        raw = soup.get_text("\n")
    return raw.strip(), suffix.lstrip(".").upper() or "TEXT"


EMAIL_RE = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.I)
SSN_RE = re.compile(r"\b\d{3}-\d{2}-\d{4}\b")
CARD_RE = re.compile(r"\b(?:\d[ -]*?){13,19}\b")
PHONE_RE = re.compile(r"(?<!\d)(?:\+?1[ .-]?)?(?:\(?[2-9]\d{2}\)?[ .-]?)\d{3}[ .-]?\d{4}(?!\d)")
DOB_RE = re.compile(
    r"(?i)\b(?:fecha\s+de\s+nacimiento|date\s+of\s+birth|dob)\s*[:#-]?\s*"
    r"(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2})"
)


def parse_custom_redactions(raw: str) -> dict[str, str]:
    if not raw.strip():
        return {}
    try:
        value = json.loads(raw)
    except json.JSONDecodeError as exc:
        raise ValueError("redacciones_json debe ser un objeto JSON válido.") from exc
    if not isinstance(value, dict):
        raise ValueError("redacciones_json debe ser un objeto JSON: {\"texto sensible\": \"[ETIQUETA]\"}.")
    result: dict[str, str] = {}
    for key, replacement in value.items():
        key = str(key)
        replacement = str(replacement)
        if key:
            result[key] = replacement or "[REDACTADO]"
    return result


def redact_text(text: str, custom: dict[str, str]) -> tuple[str, dict[str, int]]:
    counts = {"email": 0, "telefono": 0, "ssn": 0, "tarjeta": 0, "fecha_nacimiento": 0, "personalizado": 0}

    def replace(pattern: re.Pattern[str], label: str, key: str, value: str) -> str:
        def repl(match: re.Match[str]) -> str:
            counts[key] += 1
            return value
        return pattern.sub(repl, label)

    redacted = replace(EMAIL_RE, text, "email", "[EMAIL]")
    redacted = replace(PHONE_RE, redacted, "telefono", "[TELEFONO]")
    redacted = replace(SSN_RE, redacted, "ssn", "[SSN]")
    redacted = replace(CARD_RE, redacted, "tarjeta", "[TARJETA]")
    redacted = replace(DOB_RE, redacted, "fecha_nacimiento", "[FECHA_DE_NACIMIENTO]" )

    for sensitive, replacement in sorted(custom.items(), key=lambda item: len(item[0]), reverse=True):
        occurrences = redacted.count(sensitive)
        if occurrences:
            redacted = redacted.replace(sensitive, replacement)
            counts["personalizado"] += occurrences

    return redacted, counts


def list_local_documents() -> dict[str, Any]:
    root = document_root()
    files = [
        {
            "nombre": path.name,
            "ruta_relativa": str(path.relative_to(root)),
            "extension": path.suffix.lower(),
            "tamano_bytes": path.stat().st_size,
        }
        for path in sorted(root.rglob("*"))
        if path.is_file() and path.suffix.lower() in ALLOWED_EXTENSIONS and path.stat().st_size <= MAX_FILE_BYTES
    ]
    return {
        "directorio": str(root),
        "documentos": files,
        "nota": "Solo se listan nombres y metadatos; el contenido no se devuelve.",
    }


def document_fingerprint(ruta: str) -> dict[str, Any]:
    path = safe_path(ruta)
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return {
        "nombre": path.name,
        "tamano_bytes": path.stat().st_size,
        "sha256": digest.hexdigest(),
        "contenido_devuelto": False,
    }


def anonymize_document(ruta: str, redacciones_json: str = "", max_caracteres: int = MAX_OUTPUT_CHARS) -> dict[str, Any]:
    path = safe_path(ruta)
    text, document_type = extract_text(path)
    if not text:
        return {"ok": False, "error": "No se pudo extraer texto del documento. No se genera contenido sustitutivo."}

    custom = parse_custom_redactions(redacciones_json)
    sanitized, counts = redact_text(text, custom)
    try:
        limit = max(1000, min(int(max_caracteres), MAX_OUTPUT_CHARS))
    except (TypeError, ValueError):
        limit = MAX_OUTPUT_CHARS

    truncated = len(sanitized) > limit
    if truncated:
        sanitized = sanitized[:limit] + "\n\n[CONTENIDO TRUNCADO LOCALMENTE]"

    return {
        "ok": True,
        "nombre": path.name,
        "tipo_documento": document_type,
        "texto_preparado_para_ia": sanitized,
        "redacciones": counts,
        "truncado": truncated,
        "original_devuelto": False,
        "redaccion_local_deterministica": True,
        "advertencia": "Revisa el resultado antes de enviarlo a una IA externa. La detección automática no garantiza identificar todos los datos personales o confidenciales.",
    }


def redact_document_copy(ruta: str, destino: str = "", redacciones_json: str = "") -> dict[str, Any]:
    source = safe_path(ruta)
    custom = parse_custom_redactions(redacciones_json)
    text, document_type = extract_text(source)
    sanitized, counts = redact_text(text, custom)

    if destino.strip():
        target = Path(destino).expanduser()
        if not target.is_absolute():
            target = document_root() / target
    else:
        target = source.with_name(f"{source.stem}.anonimizado.txt")

    target = target.resolve(strict=False)
    root = document_root()
    try:
        target.relative_to(root)
    except ValueError as exc:
        raise ValueError(f"La copia también debe permanecer dentro de VELUM_DOCUMENT_ROOT: {root}") from exc
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(sanitized, encoding="utf-8")

    return {
        "ok": True,
        "archivo_creado": str(target),
        "tipo_origen": document_type,
        "redacciones": counts,
        "original_modificado": False,
        "contenido_original_devuelto": False,
    }


def privacy_status() -> dict[str, Any]:
    return {
        "modo": "local-first",
        "transporte": "stdio",
        "abre_puerto": False,
        "sube_documentos": False,
        "llamadas_http_desde_las_herramientas_privadas": False,
        "almacenamiento_remoto": False,
        "directorio_documentos": str(document_root()),
        "tipos": sorted(ALLOWED_EXTENSIONS),
        "limite_archivo_mb": MAX_FILE_BYTES // (1024 * 1024),
        "importante": "Si una herramienta devuelve texto a ChatGPT, Claude u otra IA, ese texto sí sale del equipo hacia ese servicio. Use preparar_documento_para_ia para enviar solo el resultado sanitizado.",
    }
